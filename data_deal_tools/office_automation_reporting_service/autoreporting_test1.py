##学习用
def main():
    # 从万得数据库读取数据，保存可视化结果图片
    def getDataAndSavePic():
        db_engine = create_engine('oracle://lianghua:lianghua@172.16.10.101:1521/orcl', echo=False)
        DB_Session = sessionmaker(bind=db_engine)
        session = DB_Session()
        s = ("select OB_OBJECT_NAME_1090 AS STOCKNAME,"#股票名称
            "       F5_1090    AS TRDMARKETNAME,               "#交易所名称
            "       F6_1090  AS PLATENAME,                  "#版块名称
            "       F17_1090 AS TIMETOMARKET,               "#上市时间
            "       F16_1090 AS STOCKCODE                   "#股票代码
            "from wind.tb_object_1090 t where t.f4_1090='A' ")#只取A股数据
        selectsql = text(s)
        result = session.execute(selectsql)  # 执行查询语句
        df_result = pd.DataFrame(result.fetchall())
        df_result.columns = ['STOCKNAME', 'TRDMARKETNAME', 'PLATENAME', 'TIMETOMARKET','STOCKCODE']  # 列重命名
        df_result = df_result.set_index('STOCKCODE')
        session.close()
        pie_file_path = r'd:\temp\pie.png'  #饼状图图片地址
        bar_file_path = r'd:\temp\bar.png' #柱状图图片地址
        #绘制饼状图，分别计算主板、中小板和创业板股票的数量占比
        (
        df_result.groupby('PLATENAME')
        .count()
        .plot.pie(y='STOCKNAME',figsize=(6, 6),autopct='%.2f')
        )
        plt.savefig(pie_file_path)
        #绘制柱状图，获取不同年份上市股票的数量
        df_result['YEARTOMARKET']=df_result['TIMETOMARKET'].map(lambda x:None if x is None else x[0:4])
        (
        df_result.groupby('YEARTOMARKET')
        .count()
        .plot.bar(y='STOCKNAME',figsize=(8, 6))
        )
        plt.savefig(bar_file_path)
        return (df_result,pie_file_path,bar_file_path)

    # 自动生成WORD文件，定义文档模板
    from docx import Document
    from docx.shared import Inches

    def gen_docfile(df,pie_file_path,bar_file_path,doc_file_path):
        '''
        :param df_result: 数据记录，用于表格显示
        :param pie_file_path: 饼图文件显示
        :param bar_file_path: 柱状图文件显示
        :param doc_file_path: 需要保存的WORK文件路径
        :return: 无返回值
        '''
        # 新建一个文档
        document = Document()
        document.add_heading(u' 自动分析报告生成 ', 0)
        # 添加一个段落
        p = document.add_paragraph(u'python-doc模块是一个非常实用的用于自动生成报告的文档，可以自动根据读取的数据生成')
        p.add_run(u'图片').bold = True
        p.add_run(u' 和 ')
        p.add_run(u'表格').italic = True
        document.add_paragraph(u'python-doc模块可以用于：')
        #无序列表项
        document.add_paragraph(
            u'根据程序计算动态结果替换动态内容，如统计数字等', style='ListBullet'
        )
        document.add_paragraph(
            u'可以自动嵌入相应的图片和表格', style='ListBullet'
        )
        document.add_paragraph(
            u'支持各类样式进行调整', style='ListBullet'
        )

        document.add_paragraph(u'python-doc模块不足的地方：')
        document.add_paragraph(
            u'相对简单', style='ListNumber'
        )
        document.add_paragraph(
            u'暂不支持WORD文档模板', style='ListNumber'
        )

        document.add_heading(u'二、各板块统计', level=1)
        text=u'沪深两地的上市A股总共有%s只，其中沪市有 %s 只，深市有%s 只,各板块的数据占比如下所示'\
             %(str(df['STOCKNAME'].count()),\
               str(df[df['TRDMARKETNAME']=='上海']['STOCKNAME'].count()),\
               str(df[df['TRDMARKETNAME']=='深圳']['STOCKNAME'].count())
               )
        document.add_paragraph(text)
        # 插入图片，文件名可以作为参数传入，由之前的程序进行传入
        document.add_picture(pie_file_path, width=Inches(5.0))

        document.add_heading(u'三、上市时间统计', level=1)
        text=u'\n上市时间分布图如下所示，可以看出今明两年并不上上市的高峰期'
        document.add_paragraph(text)
        # 插入图片，文件名可以作为参数传入，由之前的程序进行传入
        document.add_picture(bar_file_path, width=Inches(5.0))


        document.add_heading(u'四、待上市新股统计', level=1)
        # 轮询上市时间为空的未上市股票,添加表格
        text=u'\n待上市股票列表如下'
        df['TIMETOMARKET']=df['TIMETOMARKET'].map(lambda x:'99991231' if x is None else x[0:4])
        df_newstock=df[df['TIMETOMARKET']=='99991231']
        print df_newstock
        #插入表格
        table = document.add_table(rows=len(df_newstock.index)+1, cols=3,style='Table Grid')
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = u'股票名称'
        hdr_cells[1].text = u'上市交易所'
        hdr_cells[2].text = u'上市板块'
        #编历DATAFRAME
        list_stockname=list(df_newstock['STOCKNAME'])
        list_TRDMARKETNAME=list(df_newstock['TRDMARKETNAME'])
        list_PLATENAME=list(df_newstock['PLATENAME'])
        for i in range(len(df_newstock.index)):
            row_cells = table.add_row().cells
            #注意这里PYTHON2的编码问题,多谢stackoverflow,程序员的圣地
            row_cells[0].text = unicode(list_stockname[i],'utf-8')
            row_cells[1].text = unicode(list_TRDMARKETNAME[i],'utf-8')
            row_cells[2].text = unicode(list_PLATENAME[i],'utf-8')

        document.add_page_break()
        document.save(doc_file_path)

