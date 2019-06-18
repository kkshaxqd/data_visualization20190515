#!/usr/bin/env python
# --coding:utf-8 --
from docx import Document
from docx.shared import Inches

document = Document()
document.add_heading('财务分析_test', 0)
str_0 = '在已有分析框架的基础上，每次分析时很多部分都是相似的，这就为自动化财务分析提供了机会，这里就进行一下尝试'
p = document.add_paragraph(str_0)
document.add_heading('公司简介', level=1)
str1 = '公司是中国白酒龙头，主要生产销售茅台酒及茅台系列酒销售收入占公司营业收入的'
document.add_paragraph(str1)
document.add_heading('财务分析', level=1)
document.add_paragraph( '基本财务指标', style='ListNumber' )
document.add_paragraph('营业收入')
document.add_picture(r'微信图片_20180905102627.jpg')
document.add_paragraph( '盈利能力分析', style='ListNumber' )
document.add_paragraph('净利率')
document.add_picture(r'微信图片_20181017130525.jpg')
document.save('test.docx')